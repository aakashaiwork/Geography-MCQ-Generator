import datetime
import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from groq import Groq
import pandas as pd
import streamlit as st
from streamlit_gsheets import GSheetsConnection

# --- STREAMLIT UI SETUP ---
st.set_page_config(
    page_title="UPSC Geography MCQ Generator",
    page_icon="🌍",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# --- HIDE STREAMLIT BRANDING, HEADER & FOOTER ---
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    header {visibility: hidden;}
    [data-testid="stHeader"] {display: none !important;}
    footer {visibility: hidden;}
    [data-testid="stFooter"] {display: none !important;}
    .stAppDeployButton {display: none !important;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# --- GEOGRAPHY SYLLABUS MAPPING ---
GEOGRAPHY_SYLLABUS = {
    "Physical Geography": [
        "Geomorphology (Endogenetic/Exogenetic forces, Crust, Plate Tectonics, Landforms)",
        "Climatology (Heat Budget, Circulation, Monsoons, Cyclones, Koppen/Thornthwaite)",
        "Oceanography (Bottom Topography, Salinity, Currents, Tides, Corals, UNCLOS)",
        "Biogeography (Soil Profiles, Degradation, Biomes, Deforestation, Gene Pools)",
        "Environmental Geography (Ecosystems, Degradation, Hazards, EIA, Legislation)",
    ],
    "Human Geography": [
        "Perspectives in Human Geography (Areal Differentiation, Dualism, Radical/Behavioral)",
        "Economic Geography (World Resources, Agriculture Typology, Industry Location)",
        "Population & Settlement (Demographic Models, Migration, Primate City, Rank-Size)",
        "Regional Planning (Growth Poles, Regional Imbalances, Sustainable Planning)",
        "Models, Theories & Laws (Malthus, Christaller, Weber, Von Thunen, Heartland/Rimland)",
    ],
    "Geography of India": [
        "Physical Setting (Relief, Drainage, Monsoons, Western Disturbances, Soil Types)",
        "Resources (Surface/Groundwater, Minerals, Energy Crises, Forests)",
        "Agriculture (Cropping Patterns, Green/White Revolution, Agro-Climatic Zones)",
        "Industry (Locational Factors: Steel, Textile, Chemicals; SEZs, Industrial Belts)",
        "Transport, Trade & Space (Networks, Ports, Trade Policy, Space Programme)",
        "Cultural Setting & Demography (Tribal Areas, Migration, Population Attributes)",
        "Settlements & Urbanization (City Morphology, Conurbations, Slums, Urban Planning)",
        "Regional Development & Planning (Five-Year Plans, Panchayati Raj, Watershed Mgmt)",
        "Political & Geopolitical Aspects (Federalism, Inter-state issues, Ocean Realm)",
        "Contemporary Issues (Landslides, River-linking, Disaster Mgmt, Land Use Changes)",
    ],
}

# --- SYSTEM PROMPT ENFORCING EXAM-STANDARD UPSC MCQS ---
SYSTEM_PROMPT = """
You are an expert UPSC (Union Public Service Commission) question paper setter with deep academic mastery of Geography.
Your primary objective is to generate examination-standard, authentic UPSC Prelims Multiple Choice Questions (MCQs) with comprehensive pedagogical explanations.

# QUESTION TYPOLOGIES TO DISTRIBUTE EVENLY
1. Multi-Statement Questions:
   - "Consider the following statements:" with 2, 3, or 4 statements.
   - Use combinations: "1 and 2 only", "2 and 3 only", "1 and 3 only", "1, 2 and 3", OR modern UPSC pairing: "Only one", "Only two", "All three", "None".
2. Assertion-Reason (A-R):
   - Standard 4 options testing causal relationships and logical deductions.
3. Match the Following:
   - List-I and List-II with 4 distinct items and 4 combination options.
4. Negative Statement Questions:
   - "Which of the statements given above is/are NOT correct?" or "Which among the following is EXCEPT...".
5. Cause and Effect / Analytical:
   - Testing mechanisms, environmental feedback loops, or policy outcomes.

# STRICT QUALITY & DIFFICULTY CRITERIA
- 20% Basic / Foundational Concept Questions
- 40% Moderate / Multi-dimensional Analytical Questions
- 40% Difficult / Nuanced Integration Questions
- Strictly NO trivial, one-word recall questions. Focus on high-level UPSC elimination standards.
- Every question must have 4 distinct, non-overlapping options (a, b, c, d).

# MANDATORY OUTPUT STRUCTURE PER QUESTION
Format each question cleanly using standard Markdown:

### Q[Number]. [Question Stem]
[Statements / Lists if applicable]
(a) [Option A]  
(b) [Option B]  
(c) [Option C]  
(d) [Option D]  

**Correct Answer:** ([Option])  
**Detailed Explanation:**
- **Core Concept & Correctness:** [Explain why the correct option is factually and conceptually right]
- **Elimination of Other Options:** [Explicitly explain why the other options/statements are incorrect]
- **Key Takeaway / Value Addition:** [Mention relevant Geomorphic laws, Atmospheric mechanisms, Articles, Reports, or Map references]

---
"""


# --- WORD DOCUMENT HELPER ---
def create_docx(text_content):
  doc = docx.Document()

  # Set 0.5 inch margins
  for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

  style = doc.styles["Normal"]
  font = style.font
  font.name = "Verdana"
  font.size = Pt(10)

  for line in text_content.split("\n"):
    line = line.strip()
    if not line:
      doc.add_paragraph()
      continue

    if line.startswith("# "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("# ", ""))
      run.font.size = Pt(14)
      run.bold = True
    elif line.startswith("## "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("## ", ""))
      run.font.size = Pt(12)
      run.bold = True
    elif line.startswith("### "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("### ", ""))
      run.font.size = Pt(11)
      run.bold = True
    elif line.startswith("- "):
      p = doc.add_paragraph(style="List Bullet")
      p.add_run(line.replace("- ", ""))
    else:
      p = doc.add_paragraph(line)

  buffer = io.BytesIO()
  doc.save(buffer)
  buffer.seek(0)
  return buffer


# --- GOOGLE SHEETS LOGGING & ANTI-DUPLICATION ---
def get_past_mcqs_from_db(topic_name):
  try:
    if "connections" in st.secrets and "gsheets" in st.secrets["connections"]:
      conn = st.connection("gsheets", type=GSheetsConnection)
      df = conn.read(ttl="0s")
      if df is not None and not df.empty and "Topic" in df.columns:
        matching_rows = df[
            df["Topic"].str.contains(topic_name[:20], case=False, na=False)
        ]
        if not matching_rows.empty and "Question_Snippet" in matching_rows.columns:
          return matching_rows["Question_Snippet"].dropna().tolist()
  except Exception:
    pass
  return []


def save_mcq_to_db(section, topic, generated_text):
  try:
    if "connections" in st.secrets and "gsheets" in st.secrets["connections"]:
      conn = st.connection("gsheets", type=GSheetsConnection)
      existing_df = conn.read(ttl="0s")

      first_q = generated_text[:500]
      if "### Q1." in generated_text:
        first_q = (
            generated_text.split("### Q1.")[1]
            .split("**Correct Answer:**")[0]
            .strip()
        )

      new_row = pd.DataFrame([{
          "Date": str(datetime.date.today()),
          "Geography_Section": section,
          "Topic": topic,
          "Question_Snippet": first_q[:400],
      }])

      if existing_df is not None and not existing_df.empty:
        updated_df = pd.concat([existing_df, new_row], ignore_index=True)
      else:
        updated_df = new_row

      conn.update(data=updated_df)
  except Exception:
    pass


# --- MAIN INTERFACE HEADER ---
st.title("🌍 UPSC Geography MCQ Generator")
st.caption("Created by Aakash Darji")

# API Key handling
groq_api_key = st.secrets.get("GROQ_API_KEY", "")
if not groq_api_key:
  groq_api_key = st.sidebar.text_input(
      "Enter Groq API Key", type="password", help="Get key from console.groq.com"
  )

# --- MAIN FORM INPUTS ---
st.subheader("📋 Select Geography Topic")

col1, col2 = st.columns([1.5, 2.5])

with col1:
  selected_section = st.selectbox(
      "1. Geography Domain", list(GEOGRAPHY_SYLLABUS.keys())
  )

with col2:
  selected_topic = st.selectbox(
      "2. Specific Syllabus Unit", GEOGRAPHY_SYLLABUS[selected_section]
  )

col3, col4, col5 = st.columns([2, 1.5, 1.2])

with col3:
  specific_subtopic = st.text_input(
      "3. Sub-theme / Keyword Focus (Optional)",
      placeholder=(
          "e.g., Coriolis force, Karst topography, Western Ghats, Jet Streams"
      ),
  )

with col4:
  selected_languages = st.multiselect(
      "4. Language(s)",
      ["English", "Gujarati", "Hindi"],
      default=["English"],
  )

with col5:
  num_questions = st.number_input(
      "5. Questions", min_value=2, max_value=15, value=5
  )

st.divider()

# --- GENERATE ACTION ---
if st.button("🚀 Generate UPSC MCQs", type="primary", use_container_width=True):
  if not groq_api_key:
    st.error(
        "Groq API Key is missing. Please add GROQ_API_KEY to Streamlit Secrets."
    )
  elif not selected_languages:
    st.warning("Please select at least one language.")
  else:
    try:
      client = Groq(api_key=groq_api_key)

      # 1. Check Anti-Duplication DB
      past_snippets = get_past_mcqs_from_db(selected_topic)
      anti_dup_prompt = ""
      if past_snippets:
        past_str = "\n- ".join(past_snippets[-8:])
        anti_dup_prompt = (
            "\n\nCRITICAL ANTI-DUPLICATION INSTRUCTION:\nDo NOT create MCQs"
            " testing the exact same facts/statements as these previously"
            " generated items:\n- "
            + past_str
        )

      # 2. Build Prompt
      subtopic_str = (
          f"Specific Focus: '{specific_subtopic.strip()}'."
          if specific_subtopic.strip()
          else ""
      )
      lang_str = ", ".join(selected_languages)

      user_prompt = f"Generate {num_questions} UPSC Prelims-level Geography MCQs. Section: '{selected_section}', Topic: '{selected_topic}'. {subtopic_str} Target Language(s): [{lang_str}]. Include a mix of Statement-based, Assertion-Reason, Match the following, and Negative type questions with exhaustive explanations and option eliminations.{anti_dup_prompt}"

      max_tokens_val = min(1200 * num_questions * len(selected_languages), 8000)

      with st.spinner(f"Generating {num_questions} MCQs via Groq..."):
        response = client.chat.completions.create(
            model="openai/gpt-oss-120b",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_completion_tokens=max_tokens_val,
        )

        generated_mcqs = response.choices[0].message.content

        # Save to Google Sheet
        save_mcq_to_db(selected_section, selected_topic, generated_mcqs)

        st.success(f"Generated {num_questions} MCQs Successfully!")

        # Download Button
        clean_name = (
            selected_topic.split("(")[0].strip().replace(" ", "_")
        )
        docx_file = create_docx(generated_mcqs)
        st.download_button(
            label=(
                "📥 Download Pre-formatted Word Document (.docx) -"
                f" {lang_str}"
            ),
            data=docx_file,
            file_name=f"UPSC_Geography_{clean_name}_{'_'.join(selected_languages)}.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            type="primary",
        )

        st.divider()
        st.markdown(generated_mcqs)

    except Exception as e:
      st.error(f"An error occurred: {str(e)}")
